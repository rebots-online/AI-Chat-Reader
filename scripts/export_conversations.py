#!/usr/bin/env python3
# Copyright (C) 2025 Robin L. M. Cheung, MBA. All rights reserved.
"""CLI tool to export conversations to various formats."""
import os
import argparse
import csv
from typing import List
from parsers.openai_parser import OpenAIParser
from parsers.anthropic_parser import AnthropicParser
from parsers.base_parser import Conversation

try:
    from docx import Document
    from openpyxl import Workbook
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
except Exception as e:  # pragma: no cover - optional dependencies
    raise SystemExit(f"Required packages are missing: {e}")


RAW_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'data', 'raw')


def find_input_files() -> dict:
    """Locate raw JSON files."""
    inputs = {}
    for name in ['claude_conversations.json', 'anthropic_conversations.json', 'example_claude_conversations.json']:
        path = os.path.join(RAW_DIR, name)
        if os.path.exists(path):
            inputs['anthropic'] = path
            break
    for name in ['openai_conversations.json', 'chatgpt_conversations.json', 'example_openai_conversations.json']:
        path = os.path.join(RAW_DIR, name)
        if os.path.exists(path):
            inputs['openai'] = path
            break
    return inputs


def load_conversations() -> List[Conversation]:
    """Parse all available conversations."""
    inputs = find_input_files()
    convs: List[Conversation] = []
    if 'openai' in inputs:
        convs.extend(OpenAIParser().parse_file(inputs['openai']))
    if 'anthropic' in inputs:
        convs.extend(AnthropicParser().parse_file(inputs['anthropic']))
    return convs


def filter_conversations(convs: List[Conversation], match: str) -> List[Conversation]:
    if not match:
        return convs
    match_lower = match.lower()
    return [c for c in convs if match_lower in c.title.lower() or match_lower in c.id.lower()]


def export_txt(conv: Conversation, out_path: str) -> None:
    with open(out_path, 'w', encoding='utf-8') as f:
        f.write(f"# {conv.title}\n\n")
        for msg in conv.messages:
            ts = msg.timestamp.strftime('%Y-%m-%d %H:%M:%S') if msg.timestamp else ''
            f.write(f"{ts} {msg.role}: {msg.content}\n\n")


def export_md(conv: Conversation, out_path: str) -> None:
    with open(out_path, 'w', encoding='utf-8') as f:
        f.write(f"# {conv.title}\n\n")
        for msg in conv.messages:
            role = msg.role.capitalize()
            ts = msg.timestamp.strftime('%Y-%m-%d %H:%M:%S') if msg.timestamp else ''
            f.write(f"**{role}** {ts}: {msg.content}\n\n")


def export_csv(conv: Conversation, out_path: str) -> None:
    with open(out_path, 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(['timestamp', 'role', 'content'])
        for msg in conv.messages:
            ts = msg.timestamp.strftime('%Y-%m-%d %H:%M:%S') if msg.timestamp else ''
            writer.writerow([ts, msg.role, msg.content])


def export_xlsx(conv: Conversation, out_path: str) -> None:
    wb = Workbook()
    ws = wb.active
    ws.append(['timestamp', 'role', 'content'])
    for msg in conv.messages:
        ts = msg.timestamp.strftime('%Y-%m-%d %H:%M:%S') if msg.timestamp else ''
        ws.append([ts, msg.role, msg.content])
    wb.save(out_path)


def export_docx(conv: Conversation, out_path: str) -> None:
    doc = Document()
    doc.add_heading(conv.title, 0)
    for msg in conv.messages:
        ts = msg.timestamp.strftime('%Y-%m-%d %H:%M:%S') if msg.timestamp else ''
        doc.add_paragraph(f"{msg.role.capitalize()} {ts}: {msg.content}")
    doc.save(out_path)


def export_pdf(conv: Conversation, out_path: str) -> None:
    c = canvas.Canvas(out_path, pagesize=letter)
    width, height = letter
    y = height - 40
    c.setFont('Helvetica-Bold', 14)
    c.drawString(40, y, conv.title)
    y -= 30
    c.setFont('Helvetica', 10)
    for msg in conv.messages:
        line = f"{msg.role.capitalize()} "
        if msg.timestamp:
            line += msg.timestamp.strftime('%Y-%m-%d %H:%M:%S') + ': '
        line += msg.content
        for chunk in [line[i:i+90] for i in range(0, len(line), 90)]:
            if y < 40:
                c.showPage()
                y = height - 40
                c.setFont('Helvetica', 10)
            c.drawString(40, y, chunk)
            y -= 14
    c.save()


EXPORTERS = {
    'txt': export_txt,
    'md': export_md,
    'csv': export_csv,
    'xlsx': export_xlsx,
    'docx': export_docx,
    'pdf': export_pdf,
}


def main() -> None:
    parser = argparse.ArgumentParser(description='Export selected conversations to various formats.')
    parser.add_argument('-m', '--match', help='Title or ID substring to match')
    parser.add_argument('-f', '--formats', default='txt', help='Comma separated list of formats (txt,md,csv,xlsx,docx,pdf)')
    parser.add_argument('-o', '--output', default='exports', help='Output directory')
    args = parser.parse_args()

    convs = filter_conversations(load_conversations(), args.match or '')
    if not convs:
        print('No conversations matched.')
        return

    os.makedirs(args.output, exist_ok=True)
    formats = [fmt.strip() for fmt in args.formats.split(',') if fmt.strip() in EXPORTERS]
    if not formats:
        print('No valid formats specified.')
        return

    for conv in convs:
        base_name = ''.join(c for c in conv.title if c.isalnum() or c in '_-')[:50] or conv.id[:8]
        for fmt in formats:
            out_path = os.path.join(args.output, f"{base_name}.{fmt}")
            EXPORTERS[fmt](conv, out_path)
            print(f"Wrote {out_path}")


if __name__ == '__main__':
    main()
